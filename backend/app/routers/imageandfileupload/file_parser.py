import io
import zlib
import struct
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional

from pypdf import PdfReader
from docx import Document
from PIL import Image
import olefile
import pytesseract

# ------------------------------------------------------------------
# HWP (HWP 5.0 OLE 파일) 텍스트 추출
# ------------------------------------------------------------------
def extract_from_hwp(file_bytes: bytes) -> str:
    f = io.BytesIO(file_bytes)
    if not olefile.isOleFile(f):
        return file_bytes.decode('utf-8', errors='ignore')

    ole = olefile.OleFileIO(f)
    dirs = ole.listdir()
    
    # BodyText 섹션 탐색
    bodytext_dirs = [d for d in dirs if d[0] == 'BodyText']
    extracted_text = []

    for d in bodytext_dirs:
        stream = ole.openstream(d)
        data = stream.read()
        
        # zlib 압축 해제 시도
        try:
            unpacked = zlib.decompress(data, -15)
        except Exception:
            unpacked = data

        # HWP 레코드 파싱 (HWPTAG_PARA_TEXT = 67)
        i = 0
        while i < len(unpacked):
            if i + 4 > len(unpacked):
                break
            header = struct.unpack('<I', unpacked[i:i+4])[0]
            rec_type = header & 0x3FF
            rec_len = (header >> 20) & 0xFFF

            if rec_type == 67: # 문단 텍스트 태그
                para_bytes = unpacked[i+4 : i+4+rec_len]
                extracted_text.append(para_bytes.decode('utf-16-le', errors='ignore'))
            
            i += 4 + rec_len

    return "\n".join(extracted_text)


# ------------------------------------------------------------------
# HWPX (ZIP + XML 기반 파일) 텍스트 추출
# ------------------------------------------------------------------
def extract_from_hwpx(file_bytes: bytes) -> str:
    f = io.BytesIO(file_bytes)
    extracted_text = []

    with zipfile.ZipFile(f) as zf:
        for name in zf.namelist():
            # Contents/section*.xml 파일에서 텍스트 노드 추출
            if name.startswith('Contents/section') and name.endswith('.xml'):
                xml_content = zf.read(name)
                root = ET.fromstring(xml_content)
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        extracted_text.append(elem.text.strip())

    return "\n".join(extracted_text)


# ------------------------------------------------------------------
# 통합 파일 텍스트 추출 메인 함수
# ------------------------------------------------------------------
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    raw_text = ""

    try:
        # 1. PDF 문서
        if ext == 'pdf':
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            text_pages = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_pages.append(text)
            raw_text = "\n".join(text_pages)

        # 2. MS Word 문서 (.docx, .doc)
        elif ext in ['docx', 'doc']:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 표(Table) 내부 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            raw_text = "\n".join(paragraphs)

        # 3. 한글 문서 (.hwp)
        elif ext == 'hwp':
            raw_text = extract_from_hwp(file_bytes)

        # 4. 한글 문서 (.hwpx)
        elif ext == 'hwpx':
            raw_text = extract_from_hwpx(file_bytes)

        # 5. 이미지 파일 (OCR)
        elif ext in ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'heic']:
            image = Image.open(io.BytesIO(file_bytes))
            # 한글 + 영어 동시 인식
            raw_text = pytesseract.image_to_string(image, lang='kor+eng')

        # 6. 일반 텍스트 문서 (.txt, .rtf)
        elif ext in ['txt', 'rtf']:
            raw_text = file_bytes.decode('utf-8', errors='ignore')

        else:
            # 기타 확장자는 utf-8 인코딩 시도
            raw_text = file_bytes.decode('utf-8', errors='ignore')

    except Exception as e:
        print(f"[File Parse Error] {filename}: {str(e)}")
        raise ValueError(f"파일 텍스트 추출 실패 ({filename}): {str(e)}")

    # 공백 정제
    cleaned_text = raw_text.strip()
    
    if not cleaned_text:
        raise ValueError(f"파일에서 읽을 수 있는 텍스트가 없습니다. (이미지 전용 PDF이거나 빈 파일일 수 있습니다.)")

    return cleaned_text